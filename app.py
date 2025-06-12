import streamlit as st
import pandas as pd
import tempfile
import os
import base64
from io import BytesIO
from menu_processor import MenuProcessor
from extract_text import extract_text
from dotenv import load_dotenv
from docx import Document
from fpdf import FPDF
from PIL import Image

# Try to import PyMuPDF, but don't fail if not available
try:
    import fitz  # PyMuPDF for PDF preview
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

load_dotenv()
OPENAI_API_KEY = os.getenv("API")


def init_page():
    st.set_page_config(page_title="Menu Processor", layout="wide")
    st.title("📋 AI Menu Card Processor")


def get_labels(lang):
    labels = {
        "English": {
            "upload": "Upload your menu file (PDF, Image, DOCX, etc.)",
            "processing": "Processing file...",
            "download": "Download as",
            "csv": "CSV",
            "pdf": "PDF",
            "docx": "Word Doc",
            "error": "❌ Error:",
            "no_preview": "No preview available for this file type.",
            "original_file": "Original File Preview",
            "editable_data": "Editable Menu Data",
        },
        "German": {
            "upload": "Lade deine Menüdatei hoch (PDF, Bild, DOCX usw.)",
            "processing": "Datei wird verarbeitet...",
            "download": "Herunterladen als",
            "csv": "CSV",
            "pdf": "PDF",
            "docx": "Word-Dokument",
            "error": "❌ Fehler:",
            "no_preview": "Für diesen Dateityp ist keine Vorschau verfügbar.",
            "original_file": "Originaldatei-Vorschau",
            "editable_data": "Bearbeitbare Menüdaten",
        }
    }
    return labels[lang]


def show_file_preview(file_bytes: bytes, filename: str):
    """Enhanced file preview function with better error handling"""
    if not file_bytes:
        st.error("No file data available for preview")
        return

    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext in [".png", ".jpg", ".jpeg", ".bmp"]:
            img = Image.open(BytesIO(file_bytes))
            st.image(img, use_container_width=True)

        elif ext == ".heic":
            st.info("HEIC files may not display properly. Consider converting to JPG.")
            try:
                img = Image.open(BytesIO(file_bytes))
                st.image(img, use_container_width=True)
            except:
                st.error("Cannot display HEIC file")

        elif ext == ".pdf":
            if PYMUPDF_AVAILABLE:
                try:
                    doc = fitz.open(stream=file_bytes, filetype="pdf") # type: ignore
                    for page_num in range(min(3, len(doc))):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5)) # type: ignore
                        img_data = pix.tobytes("png")
                        st.image(img_data, caption=f"Page {page_num + 1}", use_container_width=True)

                    if len(doc) > 3:
                        st.info(f"Showing first 3 pages of {len(doc)} total pages")
                    doc.close()
                except Exception as e:
                    st.error(f"Error displaying PDF with PyMuPDF: {str(e)}")
            else:
                try:
                    base64_pdf = base64.b64encode(file_bytes).decode('utf-8')
                    pdf_display = f'<embed src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600" type="application/pdf">'
                    st.markdown(pdf_display, unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Error displaying PDF: {str(e)}")

        elif ext == ".docx":
            try:
                doc = Document(BytesIO(file_bytes))
                text_content = []
                for paragraph in doc.paragraphs[:20]:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)

                preview_text = '\n\n'.join(text_content)
                if preview_text:
                    st.text_area("Document Preview", preview_text, height=400, key=f"docx_preview_{hash(filename)}")
                else:
                    st.info("Document appears to be empty or contains only images/tables")

                if len(doc.paragraphs) > 20:
                    st.info(f"Showing first 20 paragraphs of {len(doc.paragraphs)} total")

            except Exception as e:
                try:
                    text = extract_text(file_bytes, filename=filename)
                    st.text_area("Document Preview", text, height=400, key=f"docx_fallback_{hash(filename)}")
                except Exception as e2:
                    st.error(f"Error displaying DOCX: {str(e2)}")

        elif ext == ".txt":
            try:
                text = file_bytes.decode("utf-8")
                preview_text = text[:5000] + "\n\n... (truncated)" if len(text) > 5000 else text
                st.text_area("Text File Preview", preview_text, height=400, key=f"txt_preview_{hash(filename)}")
            except UnicodeDecodeError:
                try:
                    text = file_bytes.decode("utf-8-sig")  # Try UTF-8 with BOM
                    st.text_area("Text File Preview", text[:5000], height=400, key=f"txt_preview_utf8_{hash(filename)}")
                except Exception as e:
                    st.error(f"Error decoding text file: {str(e)}")
        else:
            # Use a generic fallback instead of a missing `labels` reference
            st.info("No preview available for this file type.")

    except Exception as e:
        st.error(f"Error displaying file: {str(e)}")
        st.info("File uploaded successfully but preview unavailable")


def download_pdf(df):
    """Generate PDF with enhanced Unicode character handling"""
    pdf = FPDF()
    pdf.add_page()
    
    # Use Arial font
    pdf.set_font('Arial', size=10)
    
    # Add title
    pdf.set_font('Arial', 'B', size=16)
    pdf.cell(0, 10, "Menu Items", ln=1, align='C')
    pdf.ln(10)
    
    # Reset to normal font
    pdf.set_font('Arial', size=10)
    
    def clean_text(text):
        """Helper function to clean text for PDF compatibility"""
        # Replace common Unicode characters with ASCII equivalents
        replacements = {
            '\u2018': "'",  # Left single quote
            '\u2019': "'",  # Right single quote
            '\u201C': '"',  # Left double quote
            '\u201D': '"',  # Right double quote
            '\u2013': '-',  # En dash
            '\u2014': '--', # Em dash
            '\u2026': '...',# Ellipsis
            '\u20AC': 'EUR',# Euro sign
            '\u00A3': 'GBP',# Pound sign
            '\u00A9': '(c)',# Copyright sign
            '\u00AE': '(R)',# Registered sign
            '\u2022': '*',  # Bullet point
            '\u00B0': 'deg',# Degree sign
        }
        
        for unicode_char, ascii_char in replacements.items():
            text = text.replace(unicode_char, ascii_char)
            
        # Remove any remaining non-Latin1 characters
        return text.encode('cp1252', 'replace').decode('cp1252')
    
    for _, row in df.iterrows():
        try:
            # Join row data and clean text
            text = ", ".join(
                f"{col}: {str(row[col])}" 
                for col in df.columns 
                if pd.notna(row[col])
            )
            text = clean_text(text)
            
            # Handle long lines by splitting
            words = text.split(', ')
            current_line = ""
            for word in words:
                if len(current_line + word) < 80:
                    current_line += word + ", "
                else:
                    pdf.cell(0, 8, current_line.rstrip(", "), ln=1)
                    current_line = word + ", "
            if current_line:
                pdf.cell(0, 8, current_line.rstrip(", "), ln=1)
            pdf.ln(2)
            
        except Exception as e:
            st.warning(f"Some characters might not display correctly in PDF: {str(e)}")
            continue
    
    try:
        return pdf.output(dest='S').encode('latin-1') # type: ignore
    except Exception as e:
        st.error(f"Error generating PDF: {str(e)}")
        # Create minimal fallback PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', size=10)
        pdf.cell(0, 10, "Menu Items (some characters were replaced)", ln=1)
        return pdf.output(dest='S').encode('latin-1') # type: ignore


def download_docx(df):
    """Generate DOCX with UTF-8 support"""
    doc = Document()
    doc.add_heading('Menu Items', 0)

    for _, row in df.iterrows():
        text = ", ".join(f"{col}: {str(row[col])}" for col in df.columns if pd.notna(row[col]))
        doc.add_paragraph(text)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def main():
    init_page()

    # Language selection
    lang = st.radio("🌐 Select Language / Sprache wählen", ["English", "German"])
    labels = get_labels(lang)

    # File uploader
    uploaded_file = st.file_uploader(
        labels["upload"],
        type=["pdf", "png", "jpg", "jpeg", "bmp", "heic", "docx", "txt"]
    )

    # Initialize processor
    processor = MenuProcessor(llm_provider="openai", api_key=OPENAI_API_KEY)

    if uploaded_file:
        left_col, right_col = st.columns([3, 2])

        with st.spinner(labels["processing"]):
            try:
                uploaded_file.seek(0)
                file_bytes = uploaded_file.read()
                items = processor.process_menu_file(file_bytes, filename=uploaded_file.name)

                df = pd.DataFrame(items)

                with left_col:
                    st.header(labels["editable_data"])
                    edited_df = st.data_editor(
                        df,
                        num_rows="dynamic",
                        key="menu_editor",
                        use_container_width=True
                    )

                    st.subheader(labels["download"])

                    # Generate download data with UTF-8 support
                    try:
                        csv_string = processor.generate_csv(edited_df.to_dict(orient="records"), "items_empty.csv")
                    except Exception as e:
                        csv_buffer = BytesIO()
                        edited_df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
                        csv_string = csv_buffer.getvalue().decode('utf-8-sig')
                        st.warning("Using fallback CSV generation")

                    docx_data = download_docx(edited_df)
                    pdf_data = download_pdf(edited_df)

                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.download_button(
                            labels["csv"],
                            data=csv_string,
                            file_name="menu.csv",
                            mime="text/csv",
                            use_container_width=True
                        )
                    with col2:
                        st.download_button(
                            labels["docx"],
                            data=docx_data,
                            file_name="menu.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    with col3:
                        st.download_button(
                            labels["pdf"],
                            data=pdf_data,
                            file_name="menu.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )

                with right_col:
                    st.header(labels["original_file"])
                    show_file_preview(file_bytes, uploaded_file.name)

            except Exception as e:
                st.error(f"{labels['error']} {str(e)}")
                with st.expander("Debug Information"):
                    st.write(f"File name: {uploaded_file.name}")
                    st.write(f"File size: {len(file_bytes) if 'file_bytes' in locals() else 'Unknown'} bytes") # type: ignore
                    st.write(f"File type: {uploaded_file.type}")

    else:
        st.info("👆 Please upload a menu file to get started")
        show_example_layout()


def show_example_layout():
    col1, col2 = st.columns([3, 2])
    with col1:
        st.subheader("📊 Editable Data (Left Side)")
        st.write("Your processed menu data will appear here as an editable table")
        st.write("Features:")
        st.write("• Edit data directly in the table")
        st.write("• Add or remove rows")
        st.write("• Download as CSV, PDF, or Word document")
    with col2:
        st.subheader("📄 File Preview (Right Side)")
        st.write("Your uploaded file preview will appear here")
        st.write("Supported formats:")
        st.write("• PDF files (with page preview)")
        st.write("• Images (PNG, JPG, JPEG, BMP)")
        st.write("• Word documents (DOCX)")
        st.write("• Text files (TXT)")

    st.markdown("---")
    st.markdown("""
    ### 📝 Instructions:
    1. **Upload** your menu file using the file uploader above
    2. **Review** the extracted data in the left column - you can edit it directly
    3. **Preview** your original file in the right column to verify accuracy
    4. **Download** the processed data in your preferred format (CSV, PDF, or Word)

    ### 🔧 Troubleshooting:
    - Make sure your OpenAI API key is properly configured
    - Ensure all required Python packages are installed
    - For best results, use clear, high-quality menu images or PDFs
    """)


if __name__ == "__main__":
    main()
